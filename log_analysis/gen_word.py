from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面边距 ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── 全局字体辅助 ─────────────────────────────────────────
def set_run_font(run, size=11, bold=False, color=None, font_name="微软雅黑"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    sizes  = {1: 16, 2: 14, 3: 12}
    colors = {1: (0x1F, 0x49, 0x7D), 2: (0x2E, 0x74, 0xB5), 3: (0x2E, 0x74, 0xB5)}
    set_run_font(run, size=sizes.get(level, 12), bold=True,
                 color=colors.get(level, (0,0,0)))
    return p

def add_paragraph(doc, text="", bold=False, size=11, indent=0, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10.5, bold=True, color=(0xFF, 0xFF, 0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)

    # 数据行
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        fill = 'DEEAF1' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_run_font(run, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

    # 列宽
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F0F0')
        pPr.append(shd)

# ════════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('验证流程提效分析报告')
set_run_font(run, size=24, bold=True, color=(0x1F, 0x49, 0x7D))

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('Verification Efficiency Improvement Analysis')
set_run_font(run, size=14, color=(0x2E, 0x74, 0xB5))

doc.add_paragraph()
doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026年2月25日')
set_run_font(run, size=12, color=(0x70, 0x70, 0x70))

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 一、背景与现状
# ════════════════════════════════════════════════════════════
add_heading(doc, '一、背景与现状', 1)

add_heading(doc, '1.1 团队规模与技术栈', 2)
add_table(doc,
    ['项目', '详情'],
    [
        ['设计规模',   '30个模块，60-80万有效代码行'],
        ['仿真工具',   'VCS + Verdi'],
        ['版本管理',   'Git/SVN（提交粒度精确到模块级）'],
        ['Spec/VPlan', '结构化表格（Excel / 在线表格）'],
    ],
    col_widths=[4, 11]
)

doc.add_paragraph()
add_heading(doc, '1.2 各阶段耗时排序', 2)
add_table(doc,
    ['排名', '阶段', '说明'],
    [
        ['1（最耗时）', '调试定位（Debug & Triage）',          '占用工程师最多时间'],
        ['2',           '回归运行与结果分析（Regression）',    '回归慢、结果处理繁'],
        ['3',           '覆盖率Closure（Coverage Closure）',   '功能覆盖率编写全为手工'],
        ['4',           '用例编写（Testcase Development）',    '相对耗时最少'],
    ],
    col_widths=[3, 7, 5]
)

doc.add_paragraph()
add_heading(doc, '1.3 Verdi工具使用现状', 2)
add_table(doc,
    ['功能', '作用', '使用状态'],
    [
        ['nTrace / nSchema',  '信号驱动/负载追踪、连接关系图', '✅ 已使用'],
        ['Signal Group (.sg)', '预定义信号组，一键加载',        '❌ 未使用'],
        ['Tcl 脚本自动化',    '批量操作波形、自动dump',        '❌ 未使用'],
        ['Assert Debug',      '断言失败时自动关联相关信号',    '❌ 未使用'],
    ],
    col_widths=[4, 7, 4]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 二、根因分析
# ════════════════════════════════════════════════════════════
add_heading(doc, '二、根因分析', 1)

add_heading(doc, '2.1 调试定位耗时根因链', 2)
add_code_block(doc, [
    '断言覆盖几乎为零（仅靠仿真末端结果比对发现问题）',
    '    ↓',
    '问题暴露点距根因距离极远',
    '    ↓',
    '需从症状出发跨模块反向追踪（最多跨30个模块）',
    '    ↓',
    '纯手工看波形 + 对照文档，极度依赖个人经验',
    '    ↓',
    '调试耗时占比最高',
])
doc.add_paragraph()
add_paragraph(doc, '波形分析卡点排序：跨模块信号追踪 > 协议时序理解 > 信号筛选', size=10.5)
add_paragraph(doc, '断言现状：协议合规、数据完整性、状态机合法性、跨模块一致性四类断言均严重缺失；自定义接口占比70%以上。', size=10.5)

doc.add_paragraph()
add_heading(doc, '2.2 回归运行耗时根因', 2)
add_code_block(doc, [
    '大量超长用例（>8h，占比高）',
    '    +',
    '回归策略依赖人工判断（无系统化分级）',
    '    ↓',
    '资源被长时间占用 + 每次跑的范围不可预期',
    '    ↓',
    '回归周期长，结果产出慢',
])
doc.add_paragraph()
add_paragraph(doc, '注：版本管理规范，可精确追踪每次变更涉及的模块（已具备Change Impact Analysis基础条件）。', size=10.5, color=(0x2E, 0x74, 0xB5))

doc.add_paragraph()
add_heading(doc, '2.3 功能覆盖率编写痛点', 2)
for item in [
    '从Spec提取覆盖场景：全靠人工阅读，容易遗漏',
    'covergroup/coverpoint代码编写：重复性劳动多',
    'cross coverage组合：不知道哪些组合有意义',
    '与VPlan的对应关系：难以追踪覆盖点完整性',
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, '2.4 Log分诊现状', 2)
for item in [
    '当前方案：简单脚本提取UVM_ERROR关键字',
    '每次回归Fail数量：<50个',
    '错误信息由各模块验证人员自行编写，格式不统一',
    '无历史经验沉淀机制，重复问题仍需人工重新分析',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 三、提效点详述
# ════════════════════════════════════════════════════════════
add_heading(doc, '三、提效点详述', 1)

# ---------- P0 ----------
add_heading(doc, 'P0 高优先级', 2)

# 提效点1
add_heading(doc, '提效点1：RAL一键生成寄存器扫描用例', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, 'UVM已内置标准扫描序列（bit_bash / hw_reset / read_write），但实例化代码仍需人工重复编写')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    '解析RAL（.ralf / SV RAL model）',
    '    ↓',
    '自动实例化标准UVM序列',
    '    ↓',
    '生成exclusion配置文件（工程师只维护排除项）',
    '    ↓',
    'RAL变更时自动重新生成，exclusion配置保留',
])
add_paragraph(doc, '预期收益：寄存器扫描用例编写工作量接近于零，工程师只需关注排除列表维护。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点2
add_heading(doc, '提效点2：状态机断言自动生成', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '状态机合法性断言缺失，FSM层面问题无法提前暴露，导致调试链路过长')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    'RTL静态分析 → 自动提取FSM（状态、转移条件）',
    '    ↓',
    '生成状态合法性断言（非法状态、非法状态跳转）',
    '    ↓',
    '插入仿真环境，问题在FSM层面提前暴露',
])
add_paragraph(doc, '预期收益：状态机相关Bug可在FSM层直接定位，大幅减少跨模块追踪深度。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点3
add_heading(doc, '提效点3：自定义接口断言半自动生成', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '自定义接口占70%以上，协议合规断言无法套用标准模板，全靠人工编写')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    'LLM解析Excel Spec中的接口协议描述',
    '    ↓',
    '提取握手规则、时序约束、数据合法性条件',
    '    ↓',
    '生成断言模板 → 工程师审核确认 → 插入TB',
])
add_paragraph(doc, '预期收益：接口协议违规可在第一现场捕获，大幅缩短调试链路。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点4
add_heading(doc, '提效点4：Change Impact Analysis 智能回归筛选', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '每次回归范围依赖人工经验决策，无效回归比例高')
add_paragraph(doc, '前提条件：已满足（Git/SVN精确变更记录）', size=10.5, bold=True, color=(0x37, 0x86, 0x40))
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    'Git变更记录 → 识别本次修改涉及的模块',
    '    ↓',
    '建立模块 → 用例覆盖映射关系（来自覆盖率数据）',
    '    ↓',
    '自动筛选受影响模块的相关用例',
    '    ↓',
    '优先运行高相关性用例，降低无效回归',
])
add_paragraph(doc, '预期收益：针对局部改动可减少60%-80%的无效回归用例。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点5
add_heading(doc, '提效点5：三级回归体系建设', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '回归策略依赖人工判断，缺乏系统化分级，问题发现滞后')
add_paragraph(doc, '分级设计：', bold=True, size=10.5)
add_table(doc,
    ['级别', '触发条件', '用例范围', '目标时长'],
    [
        ['L1 冒烟', '每次代码提交', '核心功能用例',            '<1h'],
        ['L2 夜跑', '每日定时',     'Change Impact筛选用例',  '<8h'],
        ['L3 周跑', '每周/里程碑',  '全量回归',                '不限'],
    ],
    col_widths=[3, 4, 5.5, 3]
)
doc.add_paragraph()
add_paragraph(doc, '预期收益：问题发现及时性提升，资源利用率提高，告别人工决策。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点6
add_heading(doc, '提效点6：功能覆盖率自动生成工具', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '功能覆盖率编写场景提取、代码编写、cross coverage、traceability四个环节全为手工')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    'Excel VPlan解析',
    '    ↓',
    '提取覆盖场景、信号、边界条件',
    '    ↓',
    '自动生成covergroup / coverpoint / cross coverage框架',
    '    ↓',
    '自动挂载VPlan条目ID（保证traceability）',
    '    ↓',
    '工程师审核补充特殊逻辑',
])
add_paragraph(doc, '预期收益：覆盖率编写工作量减少50%以上，遗漏率降低。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点11
add_heading(doc, '提效点11：VPlan完整性自动检查', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, 'VPlan条目与覆盖点、用例的对应关系完全靠人工review，容易产生验证漏项')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    '解析Excel VPlan中所有验证条目',
    '    ↓',
    '自动检查每个条目是否有对应coverpoint（覆盖点）',
    '    ↓',
    '自动检查每个条目是否有对应测试用例',
    '    ↓',
    '生成完整性报告：标注未覆盖、未测试的条目',
])
add_paragraph(doc, '预期收益：验证漏项在规划阶段提前暴露，避免流片前发现重大遗漏。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点12
add_heading(doc, '提效点12：TB组件与VIP跨项目复用库', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '各项目TB框架代码重复开发，公共组件（Agent、Scoreboard模板、断言库）无法沉淀复用')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
for item in [
    '建立团队级UVM组件库：通用Agent模板、Scoreboard框架、常用断言集',
    '组件库版本化管理，新项目通过配置实例化而非重新编写',
    '每个项目的可复用组件反向沉淀回库，形成正向积累',
]:
    add_bullet(doc, item)
add_paragraph(doc, '预期收益：新项目TB搭建时间减少40%-60%，质量更稳定。', size=10.5, color=(0x37, 0x86, 0x40))

doc.add_page_break()

# ---------- P1 ----------
add_heading(doc, 'P1 中优先级', 2)

# 提效点7
add_heading(doc, '提效点7：Verdi Signal Group规范化 + Tcl调试脚本库', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '调试时手动逐一找信号、重复性波形操作耗时，且严重依赖个人经验')
add_paragraph(doc, '实现方案：', bold=True, size=10.5)
for item in [
    '按模块/场景预定义Signal Group（.sg文件），一键加载相关信号组',
    '沉淀高频调试操作为Tcl脚本（自动跳转时间点、批量dump、自动截图）',
    '结合Assert Debug，断言失败时自动关联预定义信号组',
]:
    add_bullet(doc, item)
add_paragraph(doc, '预期收益：低成本，调试准备时间明显减少，新人上手速度提升。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点8
add_heading(doc, '提效点8：超长用例拆分与并行化', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '大量>8h用例长期占用farm资源，是回归慢的直接原因')
add_paragraph(doc, '实现方案：', bold=True, size=10.5)
for item in [
    '分析超长用例的场景构成，识别可拆分的独立子场景',
    '将单个超长用例拆分为多个2-4h的并行用例',
    '保持场景覆盖完整性，确保拆分前后等价',
]:
    add_bullet(doc, item)
add_paragraph(doc, '预期收益：用例时长分布趋于均匀，farm资源利用率提升。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点9
add_heading(doc, '提效点9：覆盖率驱动早停机制', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '超长仿真末期覆盖率增长极少，大量时间浪费在冗余仿真上')
add_paragraph(doc, '实现方案：', bold=True, size=10.5)
for item in [
    '监控仿真过程中覆盖率的增长曲线',
    '当覆盖率增速低于设定阈值（plateau）时，自动终止仿真',
    '记录早停时间点，供后续优化分析',
]:
    add_bullet(doc, item)
add_paragraph(doc, '预期收益：超长用例实际运行时长可缩减20%-40%。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点10
add_heading(doc, '提效点10：错误知识库 + 智能分诊系统', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '每次回归Fail需人工逐一打开log分析，重复性问题无法复用历史经验')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    '新Fail产生 → 提取错误特征',
    '（UVM_ERROR原文 + 出错模块 + 关键上下文）',
    '    ↓',
    '与知识库做相似度匹配',
    '    ↓',
    '命中 → 显示历史根因 + 解决方案 + 关联用例',
    '未命中 → 人工分析后录入知识库',
    '    ↓',
    '知识库持续积累，命中率随时间提升',
])
doc.add_paragraph()
add_paragraph(doc, '知识库数据结构：', bold=True, size=10.5)
add_table(doc,
    ['字段', '说明'],
    [
        ['错误特征',   'UVM_ERROR原文 + 模块名 + 关键信号值'],
        ['根因分类',   'DUT Bug / TB Bug / 用例问题 / 环境问题'],
        ['根因描述',   '人工填写'],
        ['解决方案',   '人工填写'],
        ['关联用例',   '哪些用例容易触发此错误'],
    ],
    col_widths=[4, 11]
)
doc.add_paragraph()
add_paragraph(doc, '预期收益：随知识库成熟，50%以上的重复性Fail可秒级定位。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点13
add_heading(doc, '提效点13：覆盖率空洞定向用例推荐', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '覆盖率空洞分析完全依赖人工阅读报告，再手动编写定向用例，效率低且方向不精准')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    '解析覆盖率报告，自动识别未覆盖的coverpoint / cross bin',
    '    ↓',
    '映射回VPlan，确定未覆盖场景的设计语义',
    '    ↓',
    '自动推荐或生成定向用例框架（约束条件 + 场景描述）',
    '    ↓',
    '工程师补充细节后即可运行',
])
add_paragraph(doc, '预期收益：覆盖率closure阶段效率提升，精准打点补覆盖，减少盲目跑回归。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点14
add_heading(doc, '提效点14：波形Dump范围智能控制', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '全量dump导致FSDB文件过大，严重拖慢Verdi加载速度，磁盘空间消耗大')
add_paragraph(doc, '实现方案：', bold=True, size=10.5)
for item in [
    '根据用例类型/模块范围，自动配置dump scope（仅dump相关层次）',
    '失败时自动扩大dump范围重跑，成功时最小化dump',
    '结合Signal Group，只dump预定义的关键信号组',
]:
    add_bullet(doc, item)
add_paragraph(doc, '预期收益：FSDB体积减少50%-80%，Verdi加载速度显著提升，磁盘压力降低。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点15
add_heading(doc, '提效点15：Seed管理系统', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, '发现Bug的种子无系统记录，问题复现依赖工程师手动记录；回归时种子选择随机，高价值种子容易丢失')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    '自动记录每次仿真的seed及其覆盖率贡献 / 是否发现Bug',
    '    ↓',
    'Bug种子永久归档，与Bug报告双向关联',
    '    ↓',
    '回归时优先选取历史高覆盖率种子',
    '    ↓',
    '支持一键用指定seed复现历史问题',
])
add_paragraph(doc, '预期收益：问题复现成本接近于零，回归种子质量提升，覆盖率爬升更快。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# 提效点16
add_heading(doc, '提效点16：Bug与用例双向关联系统', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, 'Bug修复后难以快速定位应回归哪些用例；用例Fail时无法快速关联历史相似Bug记录')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    'Bug录入时自动关联触发它的用例 + 覆盖点 + 种子',
    '    ↓',
    'Bug修复提交代码后，自动触发关联用例回归',
    '    ↓',
    '新Fail出现时，自动搜索历史相似Bug供参考',
    '    ↓',
    '形成 Bug → 用例 → 覆盖点 的完整追溯链',
])
add_paragraph(doc, '预期收益：Bug修复验证闭环自动化，避免遗漏回归，减少重复分析成本。', size=10.5, color=(0x37, 0x86, 0x40))
doc.add_paragraph()

# ---------- P2 ----------
add_heading(doc, 'P2 长期规划', 2)

# 提效点17
add_heading(doc, '提效点17：Scoreboard参考模型辅助生成', 3)
add_paragraph(doc, '解决痛点：', bold=True, size=10.5)
add_bullet(doc, 'Scoreboard参考模型是验证中知识壁垒最高、最依赖资深工程师的环节，且完全手工编写')
add_paragraph(doc, '实现思路：', bold=True, size=10.5)
add_code_block(doc, [
    '解析RTL数据路径 + Spec中的数据变换描述',
    '    ↓',
    'LLM辅助生成参考模型框架（输入输出映射关系）',
    '    ↓',
    '工程师审核补充边界条件和异常处理逻辑',
    '    ↓',
    '与TB自动集成，完成数据比对',
])
add_paragraph(doc, '预期收益：Scoreboard开发周期缩短，知识壁垒降低，中级工程师可独立完成。', size=10.5, color=(0x37, 0x86, 0x40))

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 四、实施路线图
# ════════════════════════════════════════════════════════════
add_heading(doc, '四、实施路线图', 1)
add_table(doc,
    ['阶段', '时间', '提效点', '目标'],
    [
        ['第一阶段\nQuick Win',  '1-2个月',  '①RAL一键生成寄存器扫描用例\n⑤三级回归体系建设\n⑦Verdi Signal Group + Tcl脚本库\n⑭波形Dump范围智能控制',                                             '快速见效，零成本工具优化 + 回归分级基础建立'],
        ['第二阶段\n核心工具',   '2-4个月',  '②状态机断言自动生成\n④Change Impact Analysis\n⑥功能覆盖率自动生成\n⑪VPlan完整性自动检查\n⑫TB组件与VIP复用库\n⑮Seed管理系统\n⑩错误知识库+智能分诊', '覆盖核心痛点，调试/回归/覆盖率全面提效'],
        ['第三阶段\n深度优化',   '4-6个月',  '③自定义接口断言半自动生成\n⑧超长用例拆分并行化\n⑨覆盖率驱动早停\n⑬覆盖率空洞定向用例推荐\n⑯Bug与用例双向关联',                                     '精细化优化，全流程自动化闭环'],
        ['第四阶段\n长期规划',   '6个月以上', '⑰Scoreboard参考模型辅助生成',                                                                                                                          '突破知识壁垒，AI赋能核心验证环节'],
    ],
    col_widths=[3, 2.5, 9, 4]
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 五、优先级总览
# ════════════════════════════════════════════════════════════
add_heading(doc, '五、优先级总览', 1)
add_table(doc,
    ['优先级', '编号', '提效点', '所在阶段', '核心痛点', '难度'],
    [
        ['P0', '①',  'RAL一键生成寄存器扫描用例',       'TB/用例开发', '重复劳动',    '⭐'],
        ['P0', '②',  '状态机断言自动生成',               '调试定位',    '调试链路长',  '⭐⭐'],
        ['P0', '③',  '自定义接口断言半自动生成',         '调试定位',    '调试链路长',  '⭐⭐⭐'],
        ['P0', '④',  'Change Impact Analysis',           '回归管理',    '回归慢',      '⭐⭐'],
        ['P0', '⑤',  '三级回归体系建设',                 '回归管理',    '回归慢',      '⭐'],
        ['P0', '⑥',  '功能覆盖率自动生成工具',           '覆盖率编写',  '重复劳动',    '⭐⭐'],
        ['P0', '⑪',  'VPlan完整性自动检查',              '验证规划',    '验证漏项',    '⭐⭐'],
        ['P0', '⑫',  'TB组件与VIP跨项目复用库',          'TB开发',      '重复开发',    '⭐⭐'],
        ['P1', '⑦',  'Verdi Signal Group + Tcl脚本库',  '调试定位',    '波形调试慢',  '⭐'],
        ['P1', '⑧',  '超长用例拆分与并行化',             '回归管理',    '回归慢',      '⭐⭐'],
        ['P1', '⑨',  '覆盖率驱动早停机制',               '回归管理',    '回归慢',      '⭐⭐'],
        ['P1', '⑩',  '错误知识库 + 智能分诊',            '回归分诊',    '分诊耗时',    '⭐⭐'],
        ['P1', '⑬',  '覆盖率空洞定向用例推荐',           '覆盖率Closure','覆盖效率低', '⭐⭐⭐'],
        ['P1', '⑭',  '波形Dump范围智能控制',             '仿真执行',    '磁盘/速度',   '⭐'],
        ['P1', '⑮',  'Seed管理系统',                     '回归管理',    '复现困难',    '⭐⭐'],
        ['P1', '⑯',  'Bug与用例双向关联系统',            '报告管理',    '闭环缺失',    '⭐⭐⭐'],
        ['P2', '⑰',  'Scoreboard参考模型辅助生成',       'TB开发',      '知识壁垒',    '⭐⭐⭐⭐'],
    ],
    col_widths=[1.5, 1, 5.5, 2.8, 2.8, 1.8]
)

doc.add_paragraph()
note = doc.add_paragraph()
run = note.add_run('难度说明：⭐ = 低（1-2周）  ⭐⭐ = 中（1个月）  ⭐⭐⭐ = 较高（2个月+）  ⭐⭐⭐⭐ = 高（需LLM深度集成）')
set_run_font(run, size=9, color=(0x70, 0x70, 0x70))

# 文档末尾
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = end_p.add_run('文档生成日期：2026年2月25日')
set_run_font(run, size=9, color=(0x70, 0x70, 0x70))

doc.save(r'D:\tools\log_analysis\验证流程提效分析报告.docx')
print("Done")

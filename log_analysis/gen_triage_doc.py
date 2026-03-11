from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 页面边距 ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── 辅助函数 ─────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, color=None, font_name="微软雅黑"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    sizes  = {1: 16, 2: 14, 3: 12}
    colors = {1: (0x1F, 0x49, 0x7D), 2: (0x2E, 0x74, 0xB5), 3: (0x2E, 0x74, 0xB5)}
    set_run_font(run, size=sizes.get(level, 12), bold=True,
                 color=colors.get(level, (0, 0, 0)))
    return p

def add_paragraph(doc, text="", bold=False, size=11, indent=0, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10.5, bold=True, color=(0xFF, 0xFF, 0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'DEEAF1' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_run_font(run, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        pPr.append(shd)

def add_info_box(doc, text, fill='EBF3FB', border_color='2E74B5'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=(0x1F, 0x49, 0x7D))
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

# ════════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('仿真日志分类分诊工具')
set_run_font(run, size=26, bold=True, color=(0x1F, 0x49, 0x7D))

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('Simulation Log Triage Tool — 设计方案')
set_run_font(run, size=14, color=(0x2E, 0x74, 0xB5))

doc.add_paragraph()
doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026年2月25日')
set_run_font(run, size=12, color=(0x70, 0x70, 0x70))

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 一、工具概述
# ════════════════════════════════════════════════════════════
add_heading(doc, '一、工具概述', 1)
add_paragraph(doc,
    '本工具面向芯片验证团队，用于对仿真回归产生的日志文件进行自动化分类、提取与智能匹配分诊，'
    '帮助验证工程师快速定位报错根因，减少人工逐一翻阅日志的时间消耗，并通过知识库的持续积累'
    '实现重复性报错的秒级定位。',
    size=11)

doc.add_paragraph()
add_heading(doc, '1.1 工具定位', 2)
add_table(doc,
    ['维度', '说明'],
    [
        ['目标用户',   '芯片验证工程师'],
        ['核心价值',   '日志分类 + 首错定位 + 知识库匹配 + 经验沉淀'],
        ['运行环境',   'Windows / Linux 双平台支持'],
        ['交互方式',   'Web GUI（浏览器访问，无需安装桌面依赖）'],
        ['输出形式',   'Web页面展示 + Excel报告 + HTML报告'],
    ],
    col_widths=[4, 11]
)

doc.add_paragraph()
add_heading(doc, '1.2 解决的核心痛点', 2)
for item in [
    '每次回归Fail需人工逐一打开日志查找报错，效率低',
    '重复出现的报错无历史经验可查，每次重新分析',
    '首错信息淹没在大量Warning/Error中，难以快速定位',
    '报错根因经验分散在个人脑海中，无法团队共享',
    '知识积累无闭环机制，分析结果无法反哺知识库',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 二、整体架构
# ════════════════════════════════════════════════════════════
add_heading(doc, '二、整体架构', 1)

add_heading(doc, '2.1 技术栈', 2)
add_table(doc,
    ['层次', '技术选型', '说明'],
    [
        ['后端框架', 'Python + Flask',    '轻量Web框架，跨平台，零额外配置'],
        ['日志解析', 'Python re（正则）', '标准UVM格式精准提取'],
        ['知识库读写', 'openpyxl',        'Excel文件读写，工程师直接可维护'],
        ['报告生成', 'openpyxl + Jinja2', 'Excel报告 + HTML报告双输出'],
        ['前端展示', 'HTML / CSS / JS',   '浏览器原生支持，无需额外框架'],
    ],
    col_widths=[3, 4, 8]
)

doc.add_paragraph()
add_heading(doc, '2.2 目录结构', 2)
add_code_block(doc, [
    'triage_tool/',
    '├── app.py                 # Flask主程序 & 路由',
    '├── core/',
    '│   ├── log_parser.py      # UVM日志解析模块',
    '│   ├── matcher.py         # 知识库匹配引擎',
    '│   ├── reporter.py        # Excel / HTML报告生成',
    '│   └── db_manager.py      # 知识库读写管理',
    '├── templates/',
    '│   ├── index.html         # 主页（文件上传）',
    '│   └── result.html        # 结果展示页',
    '├── static/',
    '│   └── style.css          # 样式文件',
    '├── uploads/               # 临时存储上传的日志文件',
    '├── reports/               # 导出报告存储目录',
    '├── error_db.xlsx          # 错误知识库（首次运行自动创建）',
    '└── requirements.txt       # 依赖清单',
])

doc.add_paragraph()
add_heading(doc, '2.3 核心处理流程', 2)
add_code_block(doc, [
    '用户通过浏览器上传日志文件（单个 / 多个）',
    '    ↓',
    'log_parser：正则解析，提取全部UVM报错',
    '按 UVM_WARNING / UVM_ERROR / UVM_FATAL 分类统计',
    '定位第一个报错（首错）记录其类型、ID、描述、时间戳',
    '    ↓',
    'matcher：对首错进行知识库匹配',
    'Step1 → 错误ID精确匹配',
    'Step2 → ID未命中，关键描述关键词包含匹配',
    'Step3 → 均未命中，标记"未匹配"',
    '    ↓',
    'result.html：颜色标注展示结果',
    'FATAL(红) / ERROR(橙) / WARNING(黄) / 未匹配(灰)',
    '    ↓',
    '未匹配条目：工程师在页面填写根因 → 一键写回error_db.xlsx',
    '    ↓',
    '导出 Excel报告 / HTML报告',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 三、错误知识库设计
# ════════════════════════════════════════════════════════════
add_heading(doc, '三、错误知识库设计', 1)

add_heading(doc, '3.1 存储格式', 2)
add_paragraph(doc,
    '知识库以 Excel（.xlsx）格式存储，文件名默认为 error_db.xlsx，与工具同目录。'
    '首次运行时若文件不存在，工具自动创建含表头的空白知识库。'
    '工程师可直接用Excel打开维护，也可通过工具GUI写回。',
    size=10.5)

doc.add_paragraph()
add_heading(doc, '3.2 字段定义', 2)
add_table(doc,
    ['字段名', '类型', '是否必填', '说明', '示例'],
    [
        ['错误类型',     '枚举', '是', 'UVM_WARNING / UVM_ERROR / UVM_FATAL',  'UVM_ERROR'],
        ['错误ID',       '文本', '是', 'UVM报错的[]中的ID标识，匹配第一优先级', 'PKT_LEN_ERR'],
        ['关键描述关键词','文本', '是', '错误描述中的核心词组，多个用逗号分隔',  'length mismatch,expected'],
        ['报错原因',     '文本', '是', '根因说明，工程师填写',                  'DUT数据路径截断，包长计算有误'],
        ['所属模块',     '文本', '否', '报错来源模块名',                        'mac_rx'],
        ['根因分类',     '枚举', '是', 'DUT Bug / TB Bug / 用例问题 / 环境问题','DUT Bug'],
        ['解决方案',     '文本', '是', '处理建议',                              '检查rx_len寄存器配置'],
        ['关联用例',     '文本', '否', '容易触发此报错的用例名称',              'test_mac_rx_short_pkt'],
        ['录入人',       '文本', '否', '条目录入者',                            '张三'],
        ['录入日期',     '日期', '否', '录入时间，工具自动填写',                '2026-02-25'],
    ],
    col_widths=[3, 1.8, 2, 5, 4]
)

doc.add_paragraph()
add_heading(doc, '3.3 匹配规则', 2)
add_code_block(doc, [
    'Step 1：错误ID 精确匹配',
    '  将日志中提取的错误ID与知识库"错误ID"列完全匹配',
    '  同时校验错误类型一致 → 命中则直接返回该条目',
    '',
    'Step 2：关键描述关键词 包含匹配（ID未命中时）',
    '  将日志中的错误描述文本与知识库"关键描述关键词"列比对',
    '  知识库关键词按逗号拆分，逐个检查是否包含在描述中',
    '  所有关键词均匹配则命中',
    '',
    'Step 3：均未命中',
    '  标记为"未匹配"，触发前端回写表单',
    '  提示工程师：建议将该报错条目添加至错误数据库',
])

doc.add_paragraph()
add_info_box(doc,
    '注意：匹配时忽略大小写，提升容错率。错误ID优先级高于关键词匹配，确保准确率。')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 四、模块详细设计
# ════════════════════════════════════════════════════════════
add_heading(doc, '四、模块详细设计', 1)

# 4.1 日志解析
add_heading(doc, '4.1 log_parser.py — 日志解析模块', 2)
add_paragraph(doc, '职责：', bold=True, size=10.5)
for item in [
    '接收单个或多个日志文件路径',
    '用正则表达式提取所有UVM_WARNING / UVM_ERROR / UVM_FATAL',
    '统计各类型数量，定位首错信息',
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_paragraph(doc, '标准UVM日志格式：', bold=True, size=10.5)
add_code_block(doc, [
    'UVM_ERROR @ 1500ns [PKT_LEN_ERR] mac_rx.sv(128):',
    '  packet length mismatch: expected 64, got 32',
])

doc.add_paragraph()
add_paragraph(doc, '解析正则表达式：', bold=True, size=10.5)
add_code_block(doc, [
    'pattern = r"(UVM_(?:ERROR|WARNING|FATAL))'
    r'\s+@\s+([\d.]+\w+)\s+\[(\w+)\]\s+(\S+):\s*(.*)"',
    '# 捕获组：',
    '# Group1 → 错误级别  (UVM_ERROR)',
    '# Group2 → 时间戳    (1500ns)',
    '# Group3 → 错误ID    (PKT_LEN_ERR)',
    '# Group4 → 文件位置  (mac_rx.sv(128))',
    '# Group5 → 错误描述  (packet length mismatch...)',
])

doc.add_paragraph()
add_paragraph(doc, '输出数据结构：', bold=True, size=10.5)
add_code_block(doc, [
    '{',
    '  "file": "sim_001.log",',
    '  "statistics": {',
    '    "UVM_WARNING": 5,',
    '    "UVM_ERROR":   3,',
    '    "UVM_FATAL":   1',
    '  },',
    '  "first_error": {',
    '    "level":       "UVM_FATAL",',
    '    "timestamp":   "1500ns",',
    '    "error_id":    "PKT_LEN_ERR",',
    '    "location":    "mac_rx.sv(128)",',
    '    "description": "packet length mismatch: expected 64, got 32"',
    '  },',
    '  "all_errors": [ ... ]',
    '}',
])

doc.add_paragraph()

# 4.2 匹配引擎
add_heading(doc, '4.2 matcher.py — 知识库匹配引擎', 2)
add_paragraph(doc, '职责：', bold=True, size=10.5)
for item in [
    '加载error_db.xlsx知识库到内存',
    '对首错信息执行两阶段匹配',
    '返回匹配结果或未匹配标记',
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_paragraph(doc, '匹配逻辑伪代码：', bold=True, size=10.5)
add_code_block(doc, [
    'def match(error, db):',
    '    # Step1: 错误ID精确匹配（忽略大小写）',
    '    for entry in db:',
    '        if entry.error_id.lower() == error.error_id.lower()',
    '           and entry.level == error.level:',
    '            return {"status": "matched", "entry": entry}',
    '',
    '    # Step2: 关键词包含匹配',
    '    for entry in db:',
    '        keywords = entry.keywords.split(",")',
    '        desc = error.description.lower()',
    '        if all(kw.strip().lower() in desc for kw in keywords):',
    '            return {"status": "matched", "entry": entry}',
    '',
    '    # Step3: 未匹配',
    '    return {"status": "unmatched", "entry": None}',
])

doc.add_paragraph()

# 4.3 知识库管理
add_heading(doc, '4.3 db_manager.py — 知识库管理模块', 2)
add_paragraph(doc, '职责：', bold=True, size=10.5)
for item in [
    '首次运行时自动创建含表头的空白error_db.xlsx',
    '读取知识库所有条目供匹配引擎使用',
    '接收前端回写请求，将新条目追加到Excel末尾',
    '自动填写录入日期字段',
]:
    add_bullet(doc, item)

doc.add_paragraph()

# 4.4 报告生成
add_heading(doc, '4.4 reporter.py — 报告生成模块', 2)
add_paragraph(doc, '职责：', bold=True, size=10.5)
for item in [
    '生成Excel报告：每个日志文件对应一个Sheet，最后附汇总Sheet',
    '生成HTML报告：独立文件，支持离线查看，颜色标注',
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_paragraph(doc, 'Excel报告结构：', bold=True, size=10.5)
add_table(doc,
    ['Sheet名', '内容'],
    [
        ['sim_001',  '该日志的统计信息、首错详情、匹配结果'],
        ['sim_002',  '该日志的统计信息、首错详情、匹配结果'],
        ['汇总',     '所有日志的Pass/Fail汇总、未匹配条目列表'],
    ],
    col_widths=[4, 11]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 五、GUI界面设计
# ════════════════════════════════════════════════════════════
add_heading(doc, '五、GUI界面设计', 1)

add_heading(doc, '5.1 主页（index.html）— 文件上传', 2)
add_code_block(doc, [
    '┌─────────────────────────────────────────────────┐',
    '│          仿真日志分类分诊工具                    │',
    '├─────────────────────────────────────────────────┤',
    '│  📂 上传日志文件                                 │',
    '│  ┌─────────────────────────────────────────┐   │',
    '│  │  拖拽日志文件到此处 / 点击选择文件        │   │',
    '│  │  支持 .log .txt，可多选                  │   │',
    '│  └─────────────────────────────────────────┘   │',
    '│                                                 │',
    '│  📋 知识库路径                                   │',
    '│  [ error_db.xlsx              ] [选择文件]      │',
    '│                                                 │',
    '│              [ 开始分析 ]                       │',
    '└─────────────────────────────────────────────────┘',
])

doc.add_paragraph()
add_heading(doc, '5.2 结果页（result.html）— 分析展示', 2)
add_code_block(doc, [
    '┌─────────────────────────────────────────────────────────┐',
    '│  分析结果                   [导出Excel] [导出HTML]       │',
    '├──────────┬──────────────────────────────────────────────┤',
    '│ 文件列表 │  sim_001.log                                  │',
    '│ sim_001  ├──────────────────────────────────────────────┤',
    '│ sim_002  │  错误统计                                     │',
    '│ sim_003  │  ● UVM_FATAL   : 1   ● UVM_ERROR  : 3        │',
    '│          │  ● UVM_WARNING : 5                            │',
    '│          ├──────────────────────────────────────────────┤',
    '│          │  首错信息                              [红色] │',
    '│          │  类型：UVM_FATAL   时间戳：1500ns             │',
    '│          │  错误ID：PKT_LEN_ERR                          │',
    '│          │  位置：mac_rx.sv(128)                         │',
    '│          │  描述：packet length mismatch...              │',
    '│          ├──────────────────────────────────────────────┤',
    '│          │  匹配结果                                     │',
    '│          │  ✅ 命中知识库                                 │',
    '│          │  根因：DUT数据路径截断，包长计算逻辑有误       │',
    '│          │  分类：DUT Bug                                │',
    '│          │  方案：检查rx_len寄存器配置                   │',
    '│          │  关联用例：test_mac_rx_short_pkt              │',
    '│          ├──────────────────────────────────────────────┤',
    '│          │  ❌ 未匹配 — 请填写根因后写回知识库    [灰色] │',
    '│          │  报错原因：[___________________________]      │',
    '│          │  根因分类：[DUT Bug ▼]  所属模块：[______]   │',
    '│          │  解决方案：[___________________________]      │',
    '│          │  关联用例：[___________________________]      │',
    '│          │  录入人：  [___________________________]      │',
    '│          │                        [ 写回知识库 ]         │',
    '└──────────┴──────────────────────────────────────────────┘',
])

doc.add_paragraph()
add_heading(doc, '5.3 颜色规范', 2)
add_table(doc,
    ['颜色', '含义', '触发条件'],
    [
        ['🔴 红色',   'UVM_FATAL',   '日志中存在FATAL级别报错'],
        ['🟠 橙色',   'UVM_ERROR',   '日志中存在ERROR级别报错'],
        ['🟡 黄色',   'UVM_WARNING', '日志中存在WARNING级别报错'],
        ['🟢 绿色',   '知识库命中',  '首错在知识库中找到匹配条目'],
        ['⚫ 灰色',   '未匹配',      '首错未在知识库中找到匹配条目'],
    ],
    col_widths=[3, 4, 8]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 六、输出格式规范
# ════════════════════════════════════════════════════════════
add_heading(doc, '六、输出格式规范', 1)

add_heading(doc, '6.1 Excel报告结构', 2)
add_table(doc,
    ['Sheet', '字段', '说明'],
    [
        ['各日志Sheet\n（如sim_001）', '日志文件名\n错误统计（W/E/F数量）\n首错类型/ID/时间戳/位置/描述\n匹配状态\n报错原因/根因分类/解决方案/关联用例', '每个日志文件独立一个Sheet'],
        ['汇总Sheet',                  '文件名 | FATAL数 | ERROR数 | WARNING数 | 首错ID | 匹配状态\n未匹配条目列表', '所有日志汇总一览'],
    ],
    col_widths=[3, 9, 4]
)

doc.add_paragraph()
add_heading(doc, '6.2 HTML报告', 2)
for item in [
    '独立HTML文件，可离线浏览器打开',
    '颜色标注与GUI界面保持一致',
    '每个日志文件折叠展示，默认展开首错和匹配结果',
    '文件顶部显示整体统计（总日志数、总FATAL/ERROR/WARNING数、未匹配数）',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 七、部署与使用说明
# ════════════════════════════════════════════════════════════
add_heading(doc, '七、部署与使用说明', 1)

add_heading(doc, '7.1 环境依赖', 2)
add_code_block(doc, [
    '# requirements.txt',
    'flask>=2.0',
    'openpyxl>=3.0',
    'jinja2>=3.0      # Flask已内置',
])

doc.add_paragraph()
add_heading(doc, '7.2 安装步骤', 2)
add_code_block(doc, [
    '# Step 1: 安装依赖',
    'pip install flask openpyxl',
    '',
    '# Step 2: 启动工具',
    'python app.py',
    '',
    '# Step 3: 浏览器访问',
    'http://localhost:5000',
    '',
    '# Linux服务器远程访问（指定host）',
    'python app.py --host 0.0.0.0 --port 5000',
    'http://<server_ip>:5000',
])

doc.add_paragraph()
add_heading(doc, '7.3 使用流程', 2)
add_table(doc,
    ['步骤', '操作', '说明'],
    [
        ['1', '启动工具',         'python app.py，浏览器打开 localhost:5000'],
        ['2', '上传日志',         '拖拽或点击选择一个/多个.log文件'],
        ['3', '确认知识库路径',   '默认使用同目录error_db.xlsx，可自定义路径'],
        ['4', '点击"开始分析"',   '工具自动解析、分类、匹配'],
        ['5', '查看结果',         '左侧切换日志文件，右侧查看详情'],
        ['6', '处理未匹配条目',   '填写根因信息后点击"写回知识库"'],
        ['7', '导出报告',         '点击"导出Excel"或"导出HTML"保存报告'],
    ],
    col_widths=[1.5, 4, 10]
)

doc.add_paragraph()
add_heading(doc, '7.4 知识库维护建议', 2)
for item in [
    '每次回归结束后处理未匹配条目，及时写回知识库',
    '关键描述关键词尽量填写2-3个精准词组，避免过于宽泛导致误匹配',
    '错误ID是最精准的匹配维度，鼓励在UVM报错中设置有意义的ID',
    '定期review知识库，合并重复条目，更新过时的解决方案',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 八、后续扩展方向
# ════════════════════════════════════════════════════════════
add_heading(doc, '八、后续扩展方向', 1)
add_table(doc,
    ['扩展点', '说明', '优先级'],
    [
        ['相似度智能匹配',     '引入TF-IDF或编辑距离算法，提升模糊匹配能力',          'P2'],
        ['回归自动触发',       '回归结束后自动调用工具分析，结果推送钉钉/邮件',        'P1'],
        ['知识库统计看板',     '展示知识库条目数、命中率趋势、高频报错TOP10',          'P1'],
        ['Seed关联',           '分析结果与触发该报错的仿真seed关联记录',              'P2'],
        ['Bug系统对接',        '未匹配条目一键提交Bug（Jira/内部Bug系统）',           'P2'],
        ['多项目知识库共享',   '支持跨项目知识库合并，沉淀公司级报错经验',            'P2'],
    ],
    col_widths=[4, 9, 2.5]
)

# ── 文档末尾
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = end_p.add_run('文档生成日期：2026年2月25日')
set_run_font(run, size=9, color=(0x70, 0x70, 0x70))

doc.save(r'D:\tools\log_analysis\仿真日志分类分诊工具_设计方案.docx')
print("Done")
